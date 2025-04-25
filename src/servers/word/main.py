import os
import sys
import logging
import json
import io
from pathlib import Path
from typing import Optional, Iterable

# Add project root and src to Python path
project_root = os.path.abspath(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
)
sys.path.insert(0, project_root)
sys.path.insert(0, os.path.join(project_root, "src"))

import httpx
from docx import Document
from mcp.types import (
    Resource,
    TextContent,
    Tool,
    ImageContent,
    EmbeddedResource,
    AnyUrl,
)
from mcp.server.lowlevel.helper_types import ReadResourceContents
from mcp.server import NotificationOptions, Server
from mcp.server.models import InitializationOptions

from src.utils.microsoft.util import authenticate_and_save_credentials, get_credentials

SERVICE_NAME = Path(__file__).parent.name
MICROSOFT_GRAPH_API_URL = "https://graph.microsoft.com/v1.0"
SCOPES = [
    "Files.ReadWrite",
    "Sites.ReadWrite.All",
    "offline_access",
]

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(SERVICE_NAME)


async def make_graph_api_request(
    method,
    endpoint,
    data=None,
    params=None,
    access_token=None,
    content_type=None,
    stream=False,
):
    """Make a request to the Microsoft Graph API"""
    url = f"{MICROSOFT_GRAPH_API_URL}/{endpoint}"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": content_type if content_type else "application/json",
    }

    try:
        async with httpx.AsyncClient(follow_redirects=True) as client:
            if method.lower() == "get":
                response = await client.get(
                    url, headers=headers, params=params, timeout=60.0
                )
            elif method.lower() == "post":
                if content_type and content_type != "application/json":
                    response = await client.post(
                        url, content=data, headers=headers, params=params, timeout=60.0
                    )
                else:
                    response = await client.post(
                        url, json=data, headers=headers, params=params, timeout=60.0
                    )
            elif method.lower() == "patch":
                if content_type and content_type != "application/json":
                    response = await client.patch(
                        url, content=data, headers=headers, params=params, timeout=60.0
                    )
                else:
                    response = await client.patch(
                        url, json=data, headers=headers, params=params, timeout=60.0
                    )
            elif method.lower() == "put":
                if content_type and content_type != "application/json":
                    response = await client.put(
                        url, content=data, headers=headers, params=params, timeout=60.0
                    )
                else:
                    response = await client.put(
                        url, json=data, headers=headers, params=params, timeout=60.0
                    )
            elif method.lower() == "delete":
                response = await client.delete(
                    url, headers=headers, params=params, timeout=60.0
                )
            else:
                raise ValueError(f"Unsupported HTTP method: {method}")

            response.raise_for_status()
            if response.status_code == 204:  # No content
                return {"success": True, "status_code": 204}

            if stream:
                return response

            return response.json()

    except httpx.HTTPStatusError as e:
        error_message = f"Microsoft Graph API error: {e.response.status_code}"
        try:
            error_response = e.response.json()
            if "error" in error_response:
                error_details = error_response["error"]
                error_message = f"{error_details.get('code', 'Error')}: {error_details.get('message', 'Unknown error')}"
        except Exception:
            pass

        raise ValueError(error_message)

    except httpx.RequestError as e:
        raise ValueError(f"Failed to connect to Microsoft Graph API: {str(e)}")

    except Exception as e:
        raise ValueError(f"Error communicating with Microsoft Graph API: {str(e)}")


async def is_sharepoint_storage(access_token):
    """Detect if we're using SharePoint or OneDrive storage"""
    drive_info = await make_graph_api_request(
        "get", "me/drive", access_token=access_token
    )
    return (
        drive_info.get("driveType") == "business"
        or "sharepoint" in drive_info.get("webUrl", "").lower()
    )


def create_server(user_id, api_key=None):
    """Create a new server instance for Word operations"""
    server = Server(f"{SERVICE_NAME}-server")
    server.user_id = user_id
    server.api_key = api_key

    async def get_microsoft_client():
        """Get Microsoft access token for the current user"""
        return await get_credentials(user_id, SERVICE_NAME, api_key=api_key)

    @server.list_resources()
    async def handle_list_resources(
        cursor: Optional[str] = None,
    ) -> list[Resource]:
        """List Word documents from OneDrive"""
        access_token = await get_microsoft_client()

        try:
            # Determine if we're using SharePoint or OneDrive
            is_sharepoint = await is_sharepoint_storage(access_token)

            endpoint = "me/drive/root/search(q='.docx')"
            query_params = {
                "$top": 50,
                "$select": "id,name,webUrl,lastModifiedDateTime,size"
                + (",file" if is_sharepoint else ""),
                "$orderby": "lastModifiedDateTime desc",
            }

            if cursor:
                query_params["$skiptoken"] = cursor

            result = await make_graph_api_request(
                "get", endpoint, params=query_params, access_token=access_token
            )

            resources = []

            for item in result.get("value", []):
                # For SharePoint, filter by MIME type
                if is_sharepoint:
                    if (
                        item.get("file")
                        and item.get("file", {}).get("mimeType")
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        resources.append(
                            Resource(
                                uri=f"word:///file/{item['id']}",
                                mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                name=f"{item['name']}",
                            )
                        )
                else:
                    resources.append(
                        Resource(
                            uri=f"word:///file/{item['id']}",
                            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            name=f"{item['name']}",
                        )
                    )

            return resources

        except Exception as e:
            logger.error(f"Error fetching Word resources: {str(e)}")
            return []

    @server.read_resource()
    async def handle_read_resource(uri: AnyUrl) -> Iterable[ReadResourceContents]:
        """Read a Word document from OneDrive"""
        access_token = await get_microsoft_client()
        uri_str = str(uri)

        if uri_str.startswith("word:///file/"):
            file_id = uri_str.replace("word:///file/", "")

            try:
                endpoint = f"me/drive/items/{file_id}"
                file_info = await make_graph_api_request(
                    "get", endpoint, access_token=access_token
                )

                result = {
                    "id": file_info.get("id"),
                    "name": file_info.get("name"),
                    "webUrl": file_info.get("webUrl"),
                    "lastModifiedDateTime": file_info.get("lastModifiedDateTime"),
                    "size": file_info.get("size"),
                    "createdDateTime": file_info.get("createdDateTime"),
                    "downloadUrl": file_info.get("@microsoft.graph.downloadUrl"),
                    "contentPreview": "Content preview not available in resource view.",
                }

                formatted_content = json.dumps(result, indent=2)
                return [
                    ReadResourceContents(
                        content=formatted_content, mime_type="application/json"
                    )
                ]

            except Exception as e:
                error_msg = f"Error reading Word document: {str(e)}"
                logger.error(error_msg)

                formatted_error = {
                    "id": file_id,
                    "error": error_msg,
                    "status": "error",
                    "success": False,
                }

                return [
                    ReadResourceContents(
                        content=json.dumps(formatted_error),
                        mime_type="application/json",
                    )
                ]

        raise ValueError(f"Unsupported resource URI: {uri_str}")

    @server.list_tools()
    async def handle_list_tools() -> list[Tool]:
        """List available tools for Word"""
        return [
            Tool(
                name="list_documents",
                description="List Word documents from OneDrive",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "limit": {
                            "type": "integer",
                            "description": "Maximum number of documents to return",
                            "default": 50,
                        },
                        "query": {
                            "type": "string",
                            "description": "Optional search query to filter documents (defaults to all .docx files)",
                        },
                    },
                },
            ),
            Tool(
                name="create_document",
                description="Create a new Word document in OneDrive",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "name": {
                            "type": "string",
                            "description": "Name for the new document (will add .docx extension if not included)",
                        },
                        "content": {
                            "type": "string",
                            "description": "Content to write to the document (optional, default is empty)",
                        },
                        "folder_path": {
                            "type": "string",
                            "description": "Path to folder in OneDrive (optional, defaults to root)",
                        },
                    },
                    "required": ["name"],
                },
            ),
            Tool(
                name="read_document",
                description="Read text content from a Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "ID of the Word document",
                        },
                    },
                    "required": ["file_id"],
                },
            ),
            Tool(
                name="write_document",
                description="Append content to an existing Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "ID of the Word document",
                        },
                        "content": {
                            "type": "string",
                            "description": "Content to append to the document",
                        },
                    },
                    "required": ["file_id", "content"],
                },
            ),
            Tool(
                name="search_documents",
                description="Search for Word documents by content",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "query": {
                            "type": "string",
                            "description": "Search query to find documents containing this content",
                        },
                        "limit": {
                            "type": "integer",
                            "description": "Maximum number of documents to return",
                            "default": 25,
                        },
                    },
                    "required": ["query"],
                },
            ),
            Tool(
                name="download_document",
                description="Get a download URL for a Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "ID of the Word document",
                        },
                    },
                    "required": ["file_id"],
                },
            ),
            Tool(
                name="delete_document",
                description="Delete a Word document from OneDrive",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "file_id": {
                            "type": "string",
                            "description": "ID of the Word document",
                        },
                    },
                    "required": ["file_id"],
                },
            ),
        ]

    def format_create_document_endpoint(args):
        """Format the endpoint for creating a document"""
        file_name = args.get("name", "")
        if not file_name.lower().endswith(".docx"):
            file_name = f"{file_name}.docx"

        folder_path = args.get("folder_path", "").strip("/")
        if folder_path:
            return f"me/drive/root:/{folder_path}/{file_name}:/content"
        else:
            return f"me/drive/root:/{file_name}:/content"

    async def get_document_as_bytes(response):
        """Convert document stream to bytes"""
        doc_bytes = io.BytesIO()
        if hasattr(response, "content"):
            doc_bytes.write(response.content)
        elif hasattr(response, "read"):
            doc_bytes.write(await response.read())
        doc_bytes.seek(0)
        return doc_bytes

    @server.call_tool()
    async def handle_call_tool(
        name: str, arguments: dict | None
    ) -> list[TextContent | ImageContent | EmbeddedResource]:
        """Handle tool execution requests for Word"""
        arguments = arguments or {}
        access_token = await get_microsoft_client()

        try:
            if name == "list_documents":
                # Determine if we're using SharePoint or OneDrive
                is_sharepoint = await is_sharepoint_storage(access_token)

                endpoint = "me/drive/root/search(q='.docx')"
                params = {
                    "$top": arguments.get("limit", 50) if not is_sharepoint else 100,
                    "$select": "id,name,webUrl,lastModifiedDateTime,size,createdDateTime,file",
                    "$orderby": "lastModifiedDateTime desc",
                }

                if arguments.get("query"):
                    params["search"] = arguments.get("query")

                result = await make_graph_api_request(
                    "get", endpoint, params=params, access_token=access_token
                )

                documents = []
                if is_sharepoint:
                    for item in result.get("value", []):
                        if (
                            item.get("file")
                            and item.get("file", {}).get("mimeType")
                            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        ):
                            documents.append(item)
                else:
                    documents = result.get("value", [])

                formatted_result = {"documents": []}
                if documents and len(documents) > 0:
                    first_doc = documents[0]
                    formatted_result = {
                        "document_id": first_doc.get("id", ""),
                        "documents": [
                            {
                                "id": item.get("id"),
                                "name": item.get("name"),
                                "web_url": item.get("webUrl"),
                                "last_modified": item.get("lastModifiedDateTime"),
                                "created": item.get("createdDateTime"),
                                "size": item.get("size"),
                            }
                            for item in documents
                        ],
                    }

                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            elif name == "create_document":
                file_name = arguments.get("name", "")
                content = arguments.get("content", "")
                endpoint = format_create_document_endpoint(arguments)

                # Create a proper Word document using python-docx
                doc = Document()
                if content:
                    doc.add_paragraph(content)

                # Save the document to a byte stream
                docx_bytes = io.BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)

                result = await make_graph_api_request(
                    "put",
                    endpoint,
                    data=docx_bytes.getvalue(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    access_token=access_token,
                    params={"@microsoft.graph.conflictBehavior": "rename"},
                )

                # Check if this is SharePoint or OneDrive
                is_sharepoint = "sharepoint" in result.get("webUrl", "").lower()

                formatted_result = {
                    "created_file_id": result.get("id"),
                    "name": result.get("name"),
                    "web_url": result.get("webUrl"),
                    "content": content,
                    "is_sharepoint": is_sharepoint,
                }

                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            elif name == "read_document":
                file_id = arguments.get("file_id")

                # Get document metadata
                doc_info_endpoint = f"me/drive/items/{file_id}"
                doc_info = await make_graph_api_request(
                    "get", doc_info_endpoint, access_token=access_token
                )

                # Get document content
                content_endpoint = f"me/drive/items/{file_id}/content"
                response = await make_graph_api_request(
                    "get", content_endpoint, access_token=access_token, stream=True
                )

                # Extract document content using python-docx
                document_text = ""
                try:
                    doc_bytes = await get_document_as_bytes(response)

                    # Parse with python-docx
                    doc = Document(doc_bytes)

                    # Extract text from paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text:
                            paragraphs.append(para.text)
                    document_text = "\n".join(paragraphs)
                except Exception:
                    # Fallback to raw text if docx parsing fails
                    if hasattr(response, "text") and callable(
                        getattr(response, "text")
                    ):
                        document_text = await response.text()
                    elif hasattr(response, "content"):
                        document_text = response.content.decode(
                            "utf-8", errors="replace"
                        )

                formatted_result = {
                    "file_id": doc_info.get("id"),
                    "name": doc_info.get("name"),
                    "content": document_text,
                    "size": doc_info.get("size"),
                    "last_modified": doc_info.get("lastModifiedDateTime"),
                }

                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            elif name == "write_document":
                file_id = arguments.get("file_id")
                content = arguments.get("content")

                # Get document metadata
                doc_info_endpoint = f"me/drive/items/{file_id}"
                doc_info = await make_graph_api_request(
                    "get", doc_info_endpoint, access_token=access_token
                )

                # Get current document content
                content_endpoint = f"me/drive/items/{file_id}/content"
                response = await make_graph_api_request(
                    "get", content_endpoint, access_token=access_token, stream=True
                )

                doc_bytes = await get_document_as_bytes(response)

                try:
                    # Try to load as a Word document
                    doc = Document(doc_bytes)
                    # Add new content
                    doc.add_paragraph(content)
                except Exception:
                    # If loading as a docx fails, create a new document
                    doc = Document()

                    # Try to get current content as text
                    current_content = ""
                    if hasattr(response, "text") and callable(
                        getattr(response, "text")
                    ):
                        current_content = await response.text()
                    elif hasattr(response, "content"):
                        current_content = response.content.decode(
                            "utf-8", errors="replace"
                        )

                    if current_content:
                        doc.add_paragraph(current_content)
                    doc.add_paragraph(content)

                # Save updated document
                updated_bytes = io.BytesIO()
                doc.save(updated_bytes)
                updated_bytes.seek(0)

                # Update the document
                result = await make_graph_api_request(
                    "put",
                    content_endpoint,
                    data=updated_bytes.getvalue(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    access_token=access_token,
                    params={"@microsoft.graph.conflictBehavior": "replace"},
                )

                formatted_result = {
                    "file_id": result.get("id", file_id),
                    "name": result.get("name", doc_info.get("name")),
                    "appended": True,
                    "size": result.get("size", doc_info.get("size")),
                    "content_preview": content,
                }
                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            elif name == "search_documents":
                query = arguments.get("query")
                limit = arguments.get("limit", 25)

                # Determine if we're using SharePoint or OneDrive
                is_sharepoint = await is_sharepoint_storage(access_token)

                endpoint = f"me/drive/root/search(q='{query}')"
                params = {
                    "$top": 100 if is_sharepoint else limit,
                    "$select": "id,name,webUrl,lastModifiedDateTime,size,createdDateTime,file",
                    "$orderby": "lastModifiedDateTime desc",
                }

                if not is_sharepoint:
                    params["$filter"] = (
                        "file/mimeType eq 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'"
                    )

                result = await make_graph_api_request(
                    "get", endpoint, params=params, access_token=access_token
                )

                result_items = []
                if is_sharepoint:
                    for item in result.get("value", []):
                        if (
                            item.get("file")
                            and item.get("file", {}).get("mimeType")
                            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        ):
                            result_items.append(item)
                            if len(result_items) >= limit:
                                break
                else:
                    result_items = result.get("value", [])

                formatted_result = {
                    "documents": [],
                    "file_id": result_items[0].get("id") if result_items else "",
                    "is_sharepoint": is_sharepoint,
                }

                if result_items:
                    formatted_result["documents"] = [
                        {
                            "id": item.get("id"),
                            "name": item.get("name"),
                            "web_url": item.get("webUrl"),
                            "last_modified": item.get("lastModifiedDateTime"),
                            "created": item.get("createdDateTime"),
                            "size": item.get("size"),
                        }
                        for item in result_items
                    ]

                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            elif name == "download_document":
                file_id = arguments.get("file_id")

                endpoint = f"me/drive/items/{file_id}"
                result = await make_graph_api_request(
                    "get", endpoint, access_token=access_token
                )

                download_url = result.get("@microsoft.graph.downloadUrl")
                formatted_result = {
                    "file_id": result.get("id"),
                    "name": result.get("name"),
                    "url": download_url,
                    "size": result.get("size"),
                    "web_url": result.get("webUrl"),
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                }

                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            elif name == "delete_document":
                file_id = arguments.get("file_id")

                endpoint = f"me/drive/items/{file_id}"
                result = await make_graph_api_request(
                    "delete", endpoint, access_token=access_token
                )

                formatted_result = {
                    "deleted": True,
                    "file_id": file_id,
                    "success": True,
                }

                return [
                    TextContent(
                        type="text", text=json.dumps(formatted_result, indent=2)
                    )
                ]

            else:
                return [
                    TextContent(type="text", text=f"Error: Unsupported tool: {name}")
                ]

        except Exception as e:
            return [TextContent(type="text", text=f"Error executing {name}: {str(e)}")]

    return server


server = create_server


def get_initialization_options(server_instance: Server) -> InitializationOptions:
    """Get the initialization options for the server"""
    return InitializationOptions(
        server_name=f"{SERVICE_NAME}-server",
        server_version="1.0.0",
        capabilities=server_instance.get_capabilities(
            notification_options=NotificationOptions(),
            experimental_capabilities={},
        ),
    )


# Main entry point for authentication
if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1].lower() == "auth":
        user_id = "local"
        authenticate_and_save_credentials(user_id, SERVICE_NAME, SCOPES)
    else:
        print("Usage:")
        print("  python main.py auth - Run authentication flow for a user")
        print("Note: To run the server normally, use the guMCP server framework.")
